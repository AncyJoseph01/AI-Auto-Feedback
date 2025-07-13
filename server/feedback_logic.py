import os
import re
import warnings
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from huggingface_hub import InferenceClient
from dotenv import load_dotenv

load_dotenv()

hf_token = os.getenv("HF_TOKEN")
print(f"Using Hugging Face token: {hf_token}")
# Suppress token warnings
warnings.filterwarnings("ignore", category=UserWarning, module="huggingface_hub")

# Set your Hugging Face token
os.environ["HF_TOKEN"] = hf_token

# Initialize Hugging Face client
client = InferenceClient(
    model="meta-llama/Llama-3.1-8B-Instruct",
    token=os.environ.get("HF_TOKEN")
)

def get_all_paragraphs(doc):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
        elif child.tag.endswith('sdt'):
            for para in child.findall('.//w:p', namespaces=child.nsmap):
                yield Paragraph(para, doc)

def extract_qa(filepath):
    doc = Document(filepath)
    qa_pairs = []
    questions = []
    answers = []
    in_questions_section = True

    def is_question(text):
        return bool(re.match(r'^\d+\.\d+\s+[A-Z]', text.strip()))

    def is_section_header(text):
        return text.startswith(('LO1', 'LO2', 'LO3')) or 'Understand how to' in text or 'Glossary' in text

    for para in get_all_paragraphs(doc):
        text = para.text.strip()
        if not text:
            continue
        if is_question(text):
            questions.append(text)
            in_questions_section = True
        elif is_section_header(text):
            in_questions_section = False
        elif not in_questions_section:
            answers.append(text)

    for i, q in enumerate(questions):
        answer = answers[i] if i < len(answers) else "No answer provided"
        qa_pairs.append((q, answer))

    return qa_pairs, doc

def generate_feedback_llama(q, a, name="Student"):
    prompt = f"""
You are a professional safeguarding training assessor. Write a short, supportive, and professional feedback comment for a student named {name}.

Start with the student's name followed immediately by a comma and a space (e.g., "{name}, ...").

Focus only on the positive aspects of the student's answer.
Do not mention missing information, your own emotions, or what the teacher \"loves\".
Do not include any sign-offs like \"Best regards\" or your name at the end.

Maintain a professional tone (like a teacher giving formal written feedback).
Keep it concise: 3–5 sentences.
Vary the language and sentence style to avoid repetition across multiple answers.

Here is the question and the student's answer:

Question: {q}
Answer: {a}

Now write the feedback message:
"""
    if not a or a.strip() == "No answer provided":
        return f"{name}, it looks like there’s no answer provided for this question."

    try:
        response = client.chat.completions.create(
            model="meta-llama/Llama-3.1-8B-Instruct",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=150,
            temperature=0.85
        )
        return response.choices[0].message.content.strip()
    except Exception:
        return f"{name}, we were unable to generate feedback due to a technical issue."

def add_feedback_as_comments(doc, qa_pairs, name):
    used_qnums = set()
    for q, a in qa_pairs:
        qnum_match = re.match(r'^(\d+\.\d+)\s', q)
        if not qnum_match:
            continue
        qnum = qnum_match.group(1)
        if qnum in used_qnums:
            continue
        used_qnums.add(qnum)

        feedback = generate_feedback_llama(q, a, name)
        for para in get_all_paragraphs(doc):
            text = para.text.strip()
            if text.startswith(qnum):
                try:
                    run = para.add_run(" ")
                    doc.add_comment([run, run], feedback, author='ADMIN')
                except Exception:
                    pass
                break

def generate_overall_feedback(qa_pairs, name="Student"):
    combined_qa_text = "\n\n".join([f"Question: {q}\nAnswer: {a}" for q, a in qa_pairs])

    prompt = f"""
You are a professional safeguarding training assessor. Write a short, supportive, and professional overall feedback comment for a student named {name}.

Focus only on the positive aspects of the student's submission as a whole.
Do not mention missing information or critiques.
Do not include any sign-offs like \"Best regards\" or your name at the end.

Maintain a professional tone (like a teacher giving formal written feedback).
Keep it concise: about 5-7 sentences.
At the end, include a sentence like: congratulations \"{name}, you have passed this Unit.\"

Here is the student's entire submission:

{combined_qa_text}

Now write the overall feedback message:
"""
    try:
        response = client.chat.completions.create(
            model="meta-llama/Llama-3.1-8B-Instruct",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.85
        )
        return response.choices[0].message.content.strip()
    except Exception:
        return f"{name}, we were unable to generate overall feedback due to a technical issue."

def add_overall_feedback_comment(doc, overall_feedback, name="Student"):
    paragraphs = list(get_all_paragraphs(doc))
    if paragraphs:
        first_para = paragraphs[0]
    else:
        first_para = doc.add_paragraph()

    try:
        run = first_para.add_run(" ")
        doc.add_comment([run, run], overall_feedback, author="ADMIN")
    except Exception:
        print("Failed to add overall feedback comment.")

def process_file(input_path, output_path, student_name):
    qa_pairs, doc = extract_qa(input_path)
    add_feedback_as_comments(doc, qa_pairs, name=student_name)
    overall_feedback = generate_overall_feedback(qa_pairs, name=student_name)
    add_overall_feedback_comment(doc, overall_feedback, name=student_name)
    doc.save(output_path)
    return output_path
